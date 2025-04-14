import os
import uuid
from flask import Flask, render_template, request, send_from_directory, jsonify
from werkzeug.utils import secure_filename
from pdf2docx import Converter as PDFToDocxConverter
import subprocess
import time
import platform
from docx2pdf import convert

app = Flask(__name__)

# Configuración
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['DOWNLOAD_FOLDER'] = 'downloads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max

# Asegurar que las carpetas existan
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['DOWNLOAD_FOLDER'], exist_ok=True)

# Verificar extensión permitida
def allowed_file(filename, file_type):
    allowed_extensions = {
        'pdf': ['pdf'],
        'docx': ['docx', 'doc']
    }
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions[file_type]

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert_file():
    # Verificar si hay archivo en la solicitud
    if 'file' not in request.files:
        return jsonify({'error': 'No se encontró ningún archivo'}), 400
    
    file = request.files['file']
    conversion_type = request.form.get('conversion_type', 'pdf_to_docx')
    
    # Verificar si se seleccionó un archivo
    if file.filename == '':
        return jsonify({'error': 'No se seleccionó ningún archivo'}), 400
    
    # Determinar tipo de conversión y verificar extensión
    if conversion_type == 'pdf_to_docx':
        if not allowed_file(file.filename, 'pdf'):
            return jsonify({'error': 'Solo se permiten archivos PDF para esta conversión'}), 400
        return convert_pdf_to_docx(file)
    elif conversion_type == 'docx_to_pdf':
        if not allowed_file(file.filename, 'docx'):
            return jsonify({'error': 'Solo se permiten archivos Word (.docx, .doc) para esta conversión'}), 400
        return convert_docx_to_pdf(file)
    else:
        return jsonify({'error': 'Tipo de conversión no válido'}), 400

def convert_pdf_to_docx(file):
    try:
        # Generar nombres de archivo únicos
        unique_id = str(uuid.uuid4())
        secure_name = secure_filename(file.filename)
        base_name = os.path.splitext(secure_name)[0]
        
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{unique_id}_{secure_name}")
        docx_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f"{unique_id}_{base_name}.docx")
        
        # Guardar el archivo PDF
        file.save(pdf_path)
        
        # Convertir PDF a DOCX
        cv = PDFToDocxConverter(pdf_path)
        cv.convert(docx_path)
        cv.close()
        
        # Eliminar el archivo PDF subido para ahorrar espacio
        os.remove(pdf_path)
        
        # Devolver el nombre del archivo para descarga
        download_filename = f"{base_name}.docx"
        return jsonify({
            'success': True, 
            'message': 'Conversión exitosa', 
            'filename': os.path.basename(docx_path),
            'download_name': download_filename
        })
        
    except Exception as e:
        # En caso de error durante la conversión
        return jsonify({'error': f'Error en la conversión: {str(e)}'}), 500

def convert_docx_to_pdf(file):
    try:
        # Generar nombres de archivo únicos
        unique_id = str(uuid.uuid4())
        secure_name = secure_filename(file.filename)
        base_name = os.path.splitext(secure_name)[0]

        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{unique_id}_{secure_name}")
        pdf_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f"{unique_id}_{base_name}.pdf")

        # Guardar el archivo DOCX
        file.save(docx_path)
        
        # Verificar que LibreOffice esté instalado si estamos en Linux
        if platform.system() != "Windows":
            try:
                subprocess.run(['which', 'soffice'], check=True, stdout=subprocess.PIPE)
                app.logger.info("LibreOffice está instalado")
            except subprocess.CalledProcessError:
                app.logger.error("LibreOffice no está instalado. Necesario para docx2pdf en Linux.")
                return jsonify({
                    'error': 'Se requiere LibreOffice para la conversión en este servidor.'
                }), 500
        
        # Intentar con docx2pdf primero (requiere LibreOffice en Linux o Word en Windows)
        try:
            from docx2pdf import convert as docx2pdf_convert
            app.logger.info(f"Intentando conversión con docx2pdf")
            
            # En sistemas Linux, añadimos un timeout más largo
            if platform.system() != "Windows":
                # Usar subprocess directamente para llamar a LibreOffice
                cmd = ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', 
                       app.config['DOWNLOAD_FOLDER'], docx_path]
                process = subprocess.run(cmd, check=True, timeout=60)
                
                # Renombrar el archivo si es necesario
                generated_pdf = os.path.join(app.config['DOWNLOAD_FOLDER'], 
                                          f"{unique_id}_{secure_name}.pdf")
                if os.path.exists(generated_pdf) and generated_pdf != pdf_path:
                    os.rename(generated_pdf, pdf_path)
            else:
                # En Windows usar docx2pdf normalmente
                docx2pdf_convert(docx_path, pdf_path)
            
            app.logger.info(f"Conversión con docx2pdf exitosa")
            
        except Exception as docx2pdf_error:
            app.logger.error(f"Error con docx2pdf: {str(docx2pdf_error)}")
            
            # Intentar con otro método: PyMuPDF (requiere instalación adicional)
            try:
                app.logger.info(f"Intentando conversión con pymupdf")
                import fitz  # PyMuPDF
                
                # Convertir usando la biblioteca pymupdf
                doc = fitz.open(docx_path)
                doc.save(pdf_path)
                doc.close()
                app.logger.info(f"Conversión con pymupdf exitosa")
                
            except Exception as pymupdf_error:
                app.logger.error(f"Error con pymupdf: {str(pymupdf_error)}")
                
                # Último recurso: usar python-docx y reportlab con formato mejorado
                try:
                    app.logger.info(f"Intentando conversión con python-docx y reportlab mejorada")
                    
                    import docx
                    from reportlab.lib.pagesizes import letter
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib import colors
                    
                    # Abrir el documento docx
                    doc = docx.Document(docx_path)
                    
                    # Crear un documento PDF con un tamaño de página adecuado
                    pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
                    styles = getSampleStyleSheet()
                    
                    # Añadir estilos personalizados
                    styles.add(ParagraphStyle(name='Heading1', 
                                             parent=styles['Heading1'],
                                             fontSize=18,
                                             spaceAfter=12))
                    styles.add(ParagraphStyle(name='Heading2', 
                                             parent=styles['Heading2'],
                                             fontSize=16,
                                             spaceAfter=10))
                    styles.add(ParagraphStyle(name='Heading3', 
                                             parent=styles['Heading3'],
                                             fontSize=14,
                                             spaceAfter=8))
                    
                    story = []
                    
                    # Procesar encabezados y párrafos con formato mejorado
                    for para in doc.paragraphs:
                        if not para.text:
                            continue
                            
                        # Determinar estilo basado en el nivel de encabezado
                        if para.style.name.startswith('Heading'):
                            level = para.style.name[-1]
                            style_name = f'Heading{level}' if level in '123' else 'Normal'
                        else:
                            style_name = 'Normal'
                        
                        # Verificar si hay formato en ejecución y aplicarlo
                        formatted_text = ''
                        for run in para.runs:
                            text = run.text
                            
                            # Aplicar formato básico
                            if run.bold:
                                text = f"<b>{text}</b>"
                            if run.italic:
                                text = f"<i>{text}</i>"
                            if run.underline:
                                text = f"<u>{text}</u>"
                                
                            formatted_text += text
                            
                        p = Paragraph(formatted_text, styles[style_name])
                        story.append(p)
                        story.append(Spacer(1, 6))
                    
                    # Añadir tablas si las hay
                    for table in doc.tables:
                        # Código para procesar tablas (omitido por simplicidad)
                        pass
                    
                    # Guardar el PDF
                    pdf_doc.build(story)
                    
                    app.logger.info(f"Conversión con python-docx y reportlab mejorada exitosa")
                    
                except Exception as basic_error:
                    app.logger.error(f"Error en conversión básica: {str(basic_error)}")
                    return jsonify({
                        'error': f'No se pudo convertir el documento. Por favor, intente con un archivo más simple. Detalles técnicos: {str(basic_error)}'
                    }), 500
        
        # Verificar si el archivo PDF se creó correctamente
        if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) == 0:
            return jsonify({
                'error': 'La conversión no generó un PDF válido. Por favor, intente con otro documento.'
            }), 500
            
        # Limpiar el archivo DOCX temporal
        try:
            os.remove(docx_path)
        except:
            pass
            
        return pdf_path
        
    except Exception as e:
        app.logger.error(f"Error general en convert_docx_to_pdf: {str(e)}")
        return jsonify({
            'error': f'Error en el proceso de conversión: {str(e)}'
        }), 500

        # Eliminar el archivo docx subido
        os.remove(docx_path)
        
        download_filename = f"{base_name}.pdf"
        return jsonify({
            'success': True,
            'message': 'Conversión exitosa',
            'filename': os.path.basename(pdf_path),
            'download_name': download_filename
        })
        
    except Exception as e:
        app.logger.error(f"Error general: {str(e)}")
        return jsonify({'error': f'Error en la conversión: {str(e)}'}), 500


@app.route('/download/<filename>')
def download_file(filename):
    # Configurar el nombre de descarga para el usuario (eliminando el UUID)
    original_filename = "_".join(filename.split("_")[1:])
    
    return send_from_directory(
        app.config['DOWNLOAD_FOLDER'], 
        filename, 
        as_attachment=True,
        download_name=original_filename
    )

@app.route('/cleanup', methods=['POST'])
def cleanup_file():
    filename = request.json.get('filename')
    if filename:
        try:
            file_path = os.path.join(app.config['DOWNLOAD_FOLDER'], filename)
            if os.path.exists(file_path):
                os.remove(file_path)
            return jsonify({'success': True})
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    return jsonify({'error': 'Nombre de archivo no proporcionado'}), 400

# Tarea de limpieza periódica (opcional)
def cleanup_old_files():
    """Eliminar archivos antiguos (más de 1 hora)"""
    current_time = time.time()
    for folder in [app.config['UPLOAD_FOLDER'], app.config['DOWNLOAD_FOLDER']]:
        for filename in os.listdir(folder):
            file_path = os.path.join(folder, filename)
            # Si el archivo tiene más de 1 hora
            if os.path.isfile(file_path) and current_time - os.path.getmtime(file_path) > 3600:
                os.remove(file_path)

if __name__ == '__main__':
    # Limpiar archivos antiguos al iniciar
    cleanup_old_files()
    app.run(debug=True)
    

# Al final del archivo app.py, modifica:
if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)